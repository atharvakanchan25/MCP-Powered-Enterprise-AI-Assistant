from __future__ import annotations

import io
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor


def _now() -> str:
    return datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")


def generate_pdf(title: str, query: str, answer: str, citations: list[dict], reasoning: list[str]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    heading = ParagraphStyle("heading", parent=styles["Heading1"], textColor=colors.HexColor("#1a1a2e"))
    sub = ParagraphStyle("sub", parent=styles["Normal"], textColor=colors.grey, fontSize=9)
    story = [
        Paragraph(title, heading),
        Paragraph(f"Generated: {_now()}", sub),
        Spacer(1, 0.5 * cm),
        Paragraph("Query", styles["Heading2"]),
        Paragraph(query, styles["Normal"]),
        Spacer(1, 0.3 * cm),
        Paragraph("Answer", styles["Heading2"]),
        Paragraph(answer.replace("\n", "<br/>"), styles["Normal"]),
    ]

    if reasoning:
        story += [Spacer(1, 0.3 * cm), Paragraph("Reasoning Steps", styles["Heading2"])]
        for i, step in enumerate(reasoning, 1):
            story.append(Paragraph(f"{i}. {step}", styles["Normal"]))

    if citations:
        story += [Spacer(1, 0.3 * cm), Paragraph("Sources", styles["Heading2"])]
        table_data = [["File", "Page", "Score", "Excerpt"]]
        for c in citations:
            table_data.append([
                str(c.get("filename", "")),
                str(c.get("page", "")),
                str(c.get("score", "")),
                str(c.get("excerpt", ""))[:80] + "...",
            ])
        table = Table(table_data, colWidths=[4 * cm, 2 * cm, 2 * cm, 9 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ]))
        story.append(table)

    doc.build(story)
    return buf.getvalue()


def generate_docx(title: str, query: str, answer: str, citations: list[dict], reasoning: list[str]) -> bytes:
    doc = DocxDocument()
    h = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    doc.add_paragraph(f"Generated: {_now()}").runs[0].font.size = Pt(9)
    doc.add_heading("Query", level=2)
    doc.add_paragraph(query)
    doc.add_heading("Answer", level=2)
    doc.add_paragraph(answer)

    if reasoning:
        doc.add_heading("Reasoning Steps", level=2)
        for i, step in enumerate(reasoning, 1):
            doc.add_paragraph(f"{i}. {step}", style="List Number")

    if citations:
        doc.add_heading("Sources", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["File", "Page", "Score", "Excerpt"]):
            hdr[i].text = h
        for c in citations:
            row = table.add_row().cells
            row[0].text = str(c.get("filename", ""))
            row[1].text = str(c.get("page", ""))
            row[2].text = str(c.get("score", ""))
            row[3].text = str(c.get("excerpt", ""))[:80]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
